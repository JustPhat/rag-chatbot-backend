import os
import re
from pathlib import Path
from typing import List, Dict, Any

import fitz
from docx import Document

SUPPORTED_FILE_TYPES = {
    ".txt",
    ".docx",
    ".pdf"
}


# =========================
# Clean text
# =========================
def clean_text(text: str) -> str:

    if not text:
        return ""

    text = text.replace("\x00", " ")

    text = re.sub(
        r"\{\{.*?\}\}",
        " ",
        text,
        flags=re.DOTALL
    )

    text = re.sub(
        r"[ \t]+",
        " ",
        text
    )

    text = re.sub(
        r"\n\s*\n\s*\n+",
        "\n\n",
        text
    )

    lines = [
        line.strip()
        for line in text.splitlines()
    ]

    text = "\n".join(
        line for line in lines if line
    )

    return text.strip()


# =========================
# TXT
# =========================
def read_txt_file(file_path: str) -> str:

    encodings = [
        "utf-8",
        "utf-8-sig",
        "cp1258",
        "latin-1"
    ]

    for enc in encodings:

        try:

            with open(
                file_path,
                "r",
                encoding=enc
            ) as f:

                return f.read()

        except UnicodeDecodeError:
            continue

    with open(
        file_path,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as f:

        return f.read()


# =========================
# DOCX
# =========================
def extract_text_from_docx(
    file_path: str
) -> List[Dict[str, Any]]:

    doc = Document(file_path)

    blocks = []

    # paragraphs
    for para in doc.paragraphs:

        text = para.text.strip()

        if text:
            blocks.append(text)

    # tables
    for table in doc.tables:

        for row in table.rows:

            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                blocks.append(
                    " | ".join(cells)
                )

    full_text = "\n".join(blocks)

    return [
        {
            "page": None,
            "text": clean_text(full_text)
        }
    ]


# =========================
# PDF
# =========================
def extract_text_from_pdf(
    file_path: str
) -> List[Dict[str, Any]]:

    pages = []

    with fitz.open(file_path) as doc:

        if doc.needs_pass:
            raise ValueError(
                "PDF có mật khẩu."
            )

        for page_idx, page in enumerate(doc, start=1):

            text = page.get_text("text")

            text = clean_text(text)

            if text:

                pages.append(
                    {
                        "page": page_idx,
                        "text": text
                    }
                )

    return pages


# =========================
# Main extract
# =========================
def extract_text_from_file(
    file_path: str
) -> List[Dict[str, Any]]:

    ext = Path(file_path).suffix.lower()

    if ext not in SUPPORTED_FILE_TYPES:

        raise ValueError(
            f"Unsupported file type: {ext}"
        )

    # TXT
    if ext == ".txt":

        text = read_txt_file(file_path)

        return [
            {
                "page": None,
                "text": clean_text(text)
            }
        ]

    # DOCX
    if ext == ".docx":

        return extract_text_from_docx(file_path)

    # PDF
    if ext == ".pdf":

        return extract_text_from_pdf(file_path)

    raise ValueError(
        f"Unsupported file type: {ext}"
    )